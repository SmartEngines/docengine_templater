#
#  Copyright (c) 2016-2022, Smart Engines Service LLC
#  All rights reserved.
#

import os, sys, json, subprocess
import docx
import wx

class DocengineTemplaterFrame(wx.Frame):
  '''
    Main class for Smart Document Engine Templater - a simple DOCX-templater tool
    using Smart Document Engine CLI application for document recognition
  '''

  def __init__(self, *args, **kw):
    '''
      Initializing main class and templater-related variables
    '''
    super(DocengineTemplaterFrame, self).__init__(*args, **kw)

    self.template_path = None # path to a current DOCX template file
    self.keyval = {}          # current tags map

    self.makeMenuBar()

  def setupConfiguration(self, resources_path):
    '''
      Loads templater configuration and sets up UI
    '''

    # Loading configuration file
    self.resources_path = resources_path
    self.config = None
    with open(os.path.join(resources_path, 'config.json')) as js:
      self.config = json.load(js)

    self.panel = wx.Panel(self)

    # Button which loads a DOCX template file
    self.btemplate = wx.Button(self.panel, -1, label='Open template', size=wx.Size(300, 50), name='btemplate')
    self.Bind(wx.EVT_BUTTON, self.loadTemplate, self.btemplate)

    # Button for applying tags and saving a DOCX document
    self.bsave = wx.Button(self.panel, -1, label='Save document', size=wx.Size(300, 50), name='bsave')
    self.Bind(wx.EVT_BUTTON, self.saveDocument, self.bsave)    
    self.bsave.Disable() # disabled when template is not loaded

    template_separator = wx.StaticLine(self.panel, -1)

    # Buttons for loading document images which needs to be recognized
    self.buttons = []
    for session_key, session in self.config['sessions'].items():
      button = wx.Button(self.panel, -1, label='Load %s' % session['text'], size=wx.Size(300, 50), name=session_key)
      self.Bind(wx.EVT_BUTTON, self.loadImage, button)
      self.buttons.append(button)

    control_separator = wx.StaticLine(self.panel, -1)    

    # Button for clearing current state (tags map and path to DOCX template)
    self.bclear = wx.Button(self.panel, -1, label='Clear', size=wx.Size(300, 50), name='bclear')
    self.Bind(wx.EVT_BUTTON, self.clearState, self.bclear)

    # Status log
    self.tlog = wx.TextCtrl(self.panel, -1, size = wx.Size(500, 200 + 55 * len(self.buttons)), \
                            style = wx.TE_MULTILINE | wx.TE_READONLY | wx.TE_BESTWRAP)

    hor_sizer = wx.BoxSizer(wx.HORIZONTAL)

    buttons_sizer = wx.BoxSizer(wx.VERTICAL)
    rtab_sizer = wx.BoxSizer(wx.VERTICAL)

    buttons_sizer.Add(self.btemplate, wx.SizerFlags().Border(wx.ALL, 5))
    buttons_sizer.Add(self.bsave, wx.SizerFlags().Border(wx.ALL, 5))
    buttons_sizer.Add(template_separator, wx.SizerFlags().Border(wx.ALL, 5))
    for button in self.buttons:
      buttons_sizer.Add(button, wx.SizerFlags().Border(wx.ALL, 5))
    buttons_sizer.Add(control_separator, wx.SizerFlags().Border(wx.ALL, 5))
    buttons_sizer.Add(self.bclear, wx.SizerFlags().Border(wx.ALL, 5))

    rtab_sizer.Add(self.tlog, wx.SizerFlags().Border(wx.ALL, 5))
 
    hor_sizer.Add(buttons_sizer, wx.SizerFlags().Border(wx.TOP | wx.RIGHT | wx.LEFT, 25))
    hor_sizer.Add(rtab_sizer, wx.SizerFlags().Border(wx.TOP | wx.RIGHT | wx.LEFT, 25))
    self.panel.SetSizerAndFit(hor_sizer)

  def makeMenuBar(self):
    '''
      Initializes app menu bar
    '''
    fileMenu = wx.Menu()

    templateItem = fileMenu.Append(-1, '&Open template...\tCtrl-O', 'Opens template')
    saveItem = fileMenu.Append(-1, '&Save template...\tCtrl-S', 'Saves template')
    fileMenu.AppendSeparator()
    clearItem = fileMenu.Append(-1, 'Clear', 'Clears')
    fileMenu.AppendSeparator()
    exitItem = fileMenu.Append(wx.ID_EXIT)

    helpMenu = wx.Menu()
    aboutItem = helpMenu.Append(wx.ID_ABOUT)

    menuBar = wx.MenuBar()
    menuBar.Append(fileMenu, '&File')
    menuBar.Append(helpMenu, '&Help')

    self.SetMenuBar(menuBar)

    self.Bind(wx.EVT_MENU, self.loadTemplate, templateItem)
    self.Bind(wx.EVT_MENU, self.saveDocument, saveItem)
    self.Bind(wx.EVT_MENU, self.clearState, clearItem)
    self.Bind(wx.EVT_MENU, self.OnAbout, aboutItem)
    self.Bind(wx.EVT_MENU, self.OnExit, exitItem)

  def loadTemplate(self, event):
    '''
      Loads a DOCX template file
    '''
    self.tlog.AppendText('Loading template...\n')
    with wx.FileDialog(self, "Open template file", wildcard="DOCX files (*.docx)|*.docx",
                       style=wx.FD_OPEN | wx.FD_FILE_MUST_EXIST) as fileDialog:
      if fileDialog.ShowModal() == wx.ID_CANCEL:
        return

      pathname = fileDialog.GetPath()
      try:
        new_template_path = pathname
        new_template = docx.Document(new_template_path) # trying to load document
        self.tlog.AppendText('Loaded template from %s\n' % pathname)
        self.btemplate.SetLabel('Open template\nCurrent: %s' % os.path.split(pathname)[-1])
        self.bsave.Enable()
        self.template_path = new_template_path
      except Exception as e:
        self.tlog.AppendText('Cannot open file %s: %s\n' % (pathname, str(e)))

  def clearState(self, event):
    '''
      Clears current state
    '''
    self.tlog.AppendText('Clearing state...\n')
    self.template_path = None
    self.keyval = {}
    self.btemplate.SetLabel('Open template')
    self.bsave.Disable()

  def loadImage(self, event):
    '''
      Loads and recognizes document image, updating a keyvals map
    '''
    button_name = event.GetEventObject().GetName()
    self.tlog.AppendText('Loading image of %s...\n' % self.config['sessions'][button_name]['text'])

    with wx.FileDialog(self, 'Open %s image file' % self.config['sessions'][button_name]['text'], \
                       wildcard="PNG, JPG or TIF image (*.png;*.jpg;*.jpeg;*.tif;*.tiff)|*.png;*.jpg;*.jpeg;*.tif;*.tiff", \
                       style=wx.FD_OPEN | wx.FD_FILE_MUST_EXIST) as fileDialog:
      if fileDialog.ShowModal() == wx.ID_CANCEL:
        return

      pathname = fileDialog.GetPath()
      try:
        self.tlog.AppendText('Recognizing %s...\n' % pathname)
        output = subprocess.run([
          os.path.join(self.resources_path, self.config['executable']), # path to docengine_cli executable
          pathname,                                                     # path to image to be recognized
          os.path.join(self.resources_path, self.config['bundle']),     # path to Smart Document Engine configuration bundle
          self.config['sessions'][button_name]['documents_mask']        # path to document type mask
        ], capture_output = True)

        # parsing JSON output of docengine_cli
        output_json = None
        try:
          output_json = json.loads(output.stdout)
        except Exception:
          pass

        if output_json is None:
          self.tlog.AppendText('Failed to retrieve any data.\n')
        else:
          # updating tags map
          any_fields_extracted = False
          for tag in self.config['tags'].keys():
            if tag.split(':')[0] != button_name:
              continue
            prop_name = tag.split(':')[-1]
            if prop_name not in output_json.keys():
              continue
            prop_value = output_json[prop_name]
            self.keyval[self.config['tags'][tag]] = prop_value
            self.tlog.AppendText('Extracted %s: %s\n' % (self.config['tags'][tag], prop_value))
            any_fields_extracted = True

          if not any_fields_extracted:
            self.tlog.AppendText('No fields extracted.\n')

      except Exception as e:
        self.tlog.AppendText('Cannot process file %s: %s\n' % (pathname, str(e)))

  def applyTagsToParagraph(self, paragraph):
    '''
      Applies tags in keyval map to a DOCX paragraph, saving the format of a 
      run which contains the '$' character at the beginning of a tag
    '''
    for i in range(len(paragraph.runs)):
      while '$' in paragraph.runs[i].text:
        end_index = -1
        found_key = None
        composite_text = ''
        for j in range(i, len(paragraph.runs)):
          composite_text += paragraph.runs[j].text
          for key in self.keyval.keys():
            if '${%s}' % key in composite_text:
              found_key = key
              end_index = j
              break
          if found_key is not None:
            break
        if found_key is not None:
          paragraph.runs[i].text = composite_text.replace('${%s}' % found_key, self.keyval[found_key])
          for k in range(i + 1, end_index + 1):
            paragraph.runs[k].clear()
        else:
          break

  def saveDocument(self, event):
    '''
      Applies the tags and saves a new DOCX document
    '''
    if len(self.keyval) == 0:
      self.tlog.AppendText('Nothing to apply.\n')
      return

    self.tlog.AppendText('Applying values to template file %s:\n' % self.template_path)
    for k, v in self.keyval.items():
      self.tlog.AppendText('  %s: %s\n' % (k, v))

    document = docx.Document(self.template_path)

    # applying tags to document paragraphs
    for paragraph in document.paragraphs:
      self.applyTagsToParagraph(paragraph)
    # applying tags to all paragraphs in document tables
    for table in document.tables:
      for row in table.rows:
        for cell in row.cells:
          for paragraph in cell.paragraphs:
            self.applyTagsToParagraph(paragraph)

    with wx.FileDialog(self, "Save DOCX file", wildcard="DOCX files (*.docx)|*.docx", \
                       style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as fileDialog:

      if fileDialog.ShowModal() == wx.ID_CANCEL:
        return

      pathname = fileDialog.GetPath()
      if not pathname.lower().endswith('.docx'):
        new_pathname = pathname + '.docx'
        while os.path.exists(new_pathname):
          new_pathname = new_pathname[:-5] + '-copy.docx'
        pathname = new_pathname
      try:
        document.save(pathname)
        self.tlog.AppendText('Saved to %s\n' % pathname)
      except IOError:
        self.tlog.AppendText('Cannot save to file %s\n' % pathname)

  def OnExit(self, event):
    '''
      Closes the application
    '''
    self.Close(True)

  def OnAbout(self, event):
    '''
      Displays the About message
    '''
    wx.MessageBox('Smart Document Engine Templater v0.1', 'About', wx.OK | wx.ICON_INFORMATION)

if __name__ == '__main__':
  # retrieving a directory which contains the application to deduce resources path
  if getattr(sys, 'frozen', False):
      application_path = os.path.dirname(sys.executable)
  elif __file__:
      application_path = os.path.dirname(__file__)
  resources_path = os.path.join(application_path, 'resources')

  app = wx.App()
  
  frame = DocengineTemplaterFrame(None, title = "Smart Document Engine Templater", size = wx.Size(850, 500))
  frame.setupConfiguration(resources_path)
  frame.Show()
  
  app.MainLoop()